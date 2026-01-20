"""Базовые библиотеки"""
import os
import sys
from datetime import datetime

"""Сторонние библиотеки (требуеться устоновить)"""
try:
    from PIL import Image
    from PIL.ExifTags import TAGS
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False


"""ФУНКЦИИ"""
def is_supported_file(file_path):
    """Проверяет поддерживаемые форматы."""
    if not os.path.exists(file_path):
        return False
    
    ext = os.path.splitext(file_path)[1].lower()
    return ext in ['.jpg', '.jpeg', '.png', '.pdf', '.docx']

def get_file_info(file_path):
    """Основная информация о файле."""
    if not os.path.exists(file_path):
        return None
    
    try:
        file_stats = os.stat(file_path)
        return {
            'name': os.path.basename(file_path),
            'size_bytes': file_stats.st_size,
            'size_mb': round(file_stats.st_size / (1024 * 1024), 3),
            'created': datetime.fromtimestamp(file_stats.st_ctime).strftime('%Y-%m-%d %H:%M:%S'),
            'modified': datetime.fromtimestamp(file_stats.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
            'extension': os.path.splitext(file_path)[1].lower()
        }
    except Exception as e:
        print(f"Ошибка получения информации: {e}")
        return None

def extract_image_metadata(image_path):
    """Извлекает EXIF данные из изображений."""
    metadata = {}
    
    if not HAS_PIL:
        print("Для анализа изображений установите: pip install Pillow")
        return metadata
    
    try:
        image = Image.open(image_path)
        
        """Основная информация"""
        metadata['format'] = image.format
        metadata['mode'] = image.mode
        metadata['size_px'] = image.size
        metadata['width'] = image.width
        metadata['height'] = image.height
        
        """EXIF данные"""
        try:
            exifdata = image._getexif()
            if exifdata:
                metadata['has_exif'] = True
                for tag_id, value in exifdata.items():
                    tag = TAGS.get(tag_id, tag_id)
                    if isinstance(value, bytes):
                        try:
                            value = value.decode('utf-8', errors='ignore')
                        except:
                            value = str(value)
                    metadata[tag] = value
            else:
                metadata['has_exif'] = False
        except:
            metadata['has_exif'] = False
            
    except Exception as e:
        print(f"Ошибка чтения изображения: {e}")
    
    return metadata

def extract_pdf_metadata(pdf_path):
    """Извлекает метаданные из PDF файлов."""
    metadata = {}
    
    if not HAS_PDF:
        print("Для анализа PDF установите: pip install pypdf")
        return metadata
    
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            
            """Основная инфа"""
            metadata['pages'] = len(pdf_reader.pages)
            metadata['is_encrypted'] = pdf_reader.is_encrypted
            
            """Метаданные сканируемого документа"""
            if pdf_reader.metadata:
                meta = pdf_reader.metadata
                if hasattr(meta, 'author') and meta.author:
                    metadata['author'] = meta.author
                if hasattr(meta, 'title') and meta.title:
                    metadata['title'] = meta.title
                if hasattr(meta, 'creator') and meta.creator:
                    metadata['creator'] = meta.creator
                if hasattr(meta, 'producer') and meta.producer:
                    metadata['producer'] = meta.producer
                if hasattr(meta, 'subject') and meta.subject:
                    metadata['subject'] = meta.subject
                if hasattr(meta, 'creation_date'):
                    metadata['creation_date'] = str(meta.creation_date)
                    
    except Exception as e:
        print(f"Ошибка чтения PDF: {e}")
    
    return metadata

def extract_docx_metadata(docx_path):
    """Извлекает метаданные из DOCX файлов"""
    metadata = {}
    
    if not HAS_DOCX:
        print("Для анализа DOCX установите: pip install python-docx")
        return metadata
    
    try:
        doc = docx.Document(docx_path)
        
        """Основная инфа о документе"""
        metadata['paragraphs'] = len(doc.paragraphs)
        metadata['tables'] = len(doc.tables)
        
        """Свойства документа"""
        try:
            props = doc.core_properties
            if props.author:
                metadata['author'] = props.author
            if props.title:
                metadata['title'] = props.title
            if props.subject:
                metadata['subject'] = props.subject
            if props.keywords:
                metadata['keywords'] = props.keywords
            if props.created:
                metadata['created'] = str(props.created)
            if props.modified:
                metadata['modified'] = str(props.modified)
            if props.last_modified_by:
                metadata['last_modified_by'] = props.last_modified_by
        except:
            pass
            
    except Exception as e:
        print(f"Ошибка чтения DOCX: {e}")
    
    return metadata

def check_dangerous_metadata(metadata, file_type):
    """Проверка метаданных на опасную информацию"""
    dangerous = []
    
    """Критерии разных типов файлов"""
    if file_type == 'image':
        danger_keys = ['GPSInfo', 'GPSLatitude', 'GPSLongitude', 
                       'Model', 'Make', 'SerialNumber', 
                       'DateTimeOriginal', 'Artist', 'Software']
    elif file_type in ['pdf', 'docx']:
        danger_keys = ['author', 'creator', 'last_modified_by', 
                       'company', 'email', 'internal']
    else:
        danger_keys = []
    
    for key, value in metadata.items():
        key_str = str(key)
        for danger in danger_keys:
            if danger.lower() in key_str.lower():
                dangerous.append({
                    'key': key,
                    'value': str(value)[:100],
                    'reason': f"Содержит опасный ключ: {danger}"
                })
                break
    
    return dangerous

def clean_image_metadata(image_path, output_path=None):
    """Удаляет EXIF данные из изображения"""
    if not HAS_PIL:
        print("Для очистки изображений установите: pip install Pillow")
        return False, "Pillow не установлен"
    
    try:
        """Открытие изображения"""
        image = Image.open(image_path)
        
        """Создание нового изображения без метаданных"""
        data = list(image.getdata())
        clean_image = Image.new(image.mode, image.size)
        clean_image.putdata(data)
        
        """Сохранение"""
        if output_path is None:
            filename, ext = os.path.splitext(image_path)
            output_path = f"{filename}_cleaned{ext}"
        
        clean_image.save(output_path)
        return True, output_path
        
    except Exception as e:
        return False, str(e)

def display_analysis(file_info, metadata, dangerous):
    """Отображает результаты анализа"""
    print("\n" + "|"*60)
    print("АНАЛИЗ МЕТАДАННЫХ ФАЙЛА")
    print("|"*60)
    
    """Информация о файле"""
    print(f"\nИНФОРМАЦИЯ О ФАЙЛЕ:")
    print(f"  Название: {file_info['name']}")
    print(f"  Размер: {file_info['size_mb']} МБ ({file_info['size_bytes']} байт)")
    print(f"  Тип: {file_info['extension']}")
    print(f"  Создан: {file_info['created']}")
    print(f"  Изменён: {file_info['modified']}")
    
    """Метаданные"""
    if metadata:
        print(f"\nИЗВЛЕЧЕННЫЕ МЕТАДАННЫЕ:")
        count = 0
        for key, value in metadata.items():
            if count >= 15:
                remaining = len(metadata) - count
                print(f"  ... и ещё {remaining} полей")
                break
            
            value_str = str(value)
            if len(value_str) > 60:
                value_str = value_str[:57] + "..."
            print(f"  {key}: {value_str}")
            count += 1
    else:
        print("\nМетаданные не найдены или не удалось извлечь")
    
    """Опасная информация"""
    if dangerous:
        print(f"\nОПАСНАЯ ИНФОРМАЦИЯ НАЙДЕНА:")
        for item in dangerous:
            print(f"  • {item['key']}: {item['value']}")
            print(f"    Причина: {item['reason']}")
    else:
        print("\nОпасная информация не обнаружена")
    
    print("|"*60)

def analyze_file(file_path):
    """Анализирование файла"""
    if not os.path.exists(file_path):
        print(f"Ошибка: Файл не найден - {file_path}")
        return False
    
    if not is_supported_file(file_path):
        print(f"Ошибка: Формат не поддерживается - {file_path}")
        print("Поддерживаемые форматы: .jpg, .jpeg, .png, .pdf, .docx")
        return False
    
    """Получение информации о файле"""
    file_info = get_file_info(file_path)
    if not file_info:
        return False
    
    ext = file_info['extension']
    
    """Извлечение метаданных"""
    metadata = {}
    file_type = ""
    
    if ext in ['.jpg', '.jpeg', '.png']:
        file_type = 'image'
        metadata = extract_image_metadata(file_path)
    elif ext == '.pdf':
        file_type = 'pdf'
        metadata = extract_pdf_metadata(file_path)
    elif ext == '.docx':
        file_type = 'docx'
        metadata = extract_docx_metadata(file_path)
    
    """Проверка на опасную информацию"""
    dangerous = check_dangerous_metadata(metadata, file_type)
    
    """Показывает результаты"""
    display_analysis(file_info, metadata, dangerous)
    
    return True

def scan_folder(folder_path):
    """Сканирует все файлы в папке"""
    if not os.path.exists(folder_path):
        print(f"Ошибка: Папка не найдена - {folder_path}")
        return False
    
    if not os.path.isdir(folder_path):
        print(f"Ошибка: Это не папка - {folder_path}")
        return False
    
    print(f"\nСКАНИРОВАНИЕ ПАПКИ: {folder_path}")
    print("|"*60)
    
    """Находит все поддерживаемые файлы"""
    files = []
    for item in os.listdir(folder_path):
        item_path = os.path.join(folder_path, item)
        if os.path.isfile(item_path) and is_supported_file(item_path):
            files.append(item_path)
    
    if not files:
        print("Нет поддерживаемых файлов")
        return False
    
    print(f"Найдено файлов для анализа: {len(files)}")
    
    """Анализирует каждый файл"""
    for i, file_path in enumerate(files, 1):
        print(f"\n[{i}/{len(files)}] Анализ: {os.path.basename(file_path)}")
        analyze_file(file_path)
    
    return True

def show_help():
    """Показывает справку"""
    print("\nАНАЛИЗАТОР МЕТАДАННЫХ ФАЙЛОВ")
    print("|"*60)
    print("\nКоманды:")
    print("analyze <путь_к_файлу>  - Анализ файла")
    print("scan <путь_к_папке>     - Анализ всех файлов в папке")
    print("clean <путь_к_файлу>    - Очистка метаданных (изображения)")
    
    print("\nПоддерживаемые форматы:")
    print("•Изображения: JPG, JPEG, PNG (EXIF данные)")
    print("•Документы: PDF, DOCX (метаданные)")
    
    print("\nЧто анализируется:")
    print("•Размер и даты создания файлов")
    print("•EXIF данные изображений (GPS, камера, дата съёмки)")
    print("•Метаданные документов (автор, название, даты)")
    print("•Потенциально опасная информация")
    
    print("\nПримеры использования:")
    print("python examen.py analyze photo.jpg")
    print("python examen.py scan C:\\Documents")
    print("python examen.py clean image.jpg")
    
    print("\nУстановка доп библиотек:")
    print("pip install Pillow python-docx pypdf")

def main():
    """Главная функция!"""
    if len(sys.argv) < 2:
        show_help()
        return
    
    command = sys.argv[1].lower()
    
    if command == 'help':
        show_help()
    
    elif command == 'analyze' and len(sys.argv) >= 3:
        analyze_file(sys.argv[2])
    
    elif command == 'scan' and len(sys.argv) >= 3:
        scan_folder(sys.argv[2])
    
    elif command == 'clean' and len(sys.argv) >= 3:
        file_path = sys.argv[2]
        output_path = sys.argv[3] if len(sys.argv) >= 4 else None
        success, result = clean_image_metadata(file_path, output_path)
        if success:
            print(f"Файл очищен: {result}")
        else:
            print(f"Ошибка: {result}")
    
    else:
        print("Неизвестная команда")
        show_help()


main()